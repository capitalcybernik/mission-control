#!/usr/bin/env python3
"""
Generate CMMC Level 2 Policy Document Templates in Capital Cyber Format
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# Define the 14 control families and their controls
CONTROL_FAMILIES = {
    'AC': {
        'name': 'Access Control',
        'code': 'AC',
        'total_controls': 26,
        'l1_controls': [
            ('AC.L1-3.1.1', 'Authorized Access Control'),
            ('AC.L1-3.1.2', 'Transaction & Function Control'),
            ('AC.L1-3.1.20', 'External Connections'),
            ('AC.L1-3.1.22', 'Control Public Information')
        ],
        'l2_controls': [
            ('AC.3.1.3', 'Control CUI Flow'),
            ('AC.3.1.4', 'Separate Duties'),
            ('AC.3.1.5', 'Least Privilege'),
            ('AC.3.1.6', 'Non-privileged Accounts'),
            ('AC.3.1.7', 'Prevent Non-privileged Execution'),
            ('AC.3.1.8', 'Limit Unsuccessful Logons'),
            ('AC.3.1.9', 'Privacy and Security Notices'),
            ('AC.3.1.10', 'Session Lock'),
            ('AC.3.1.11', 'Session Termination'),
            ('AC.3.1.12', 'Remote Access Control'),
            ('AC.3.1.13', 'Cryptographic Remote Access'),
            ('AC.3.1.14', 'Remote Access Points'),
            ('AC.3.1.15', 'Remote Privileged Commands'),
            ('AC.3.1.16', 'Wireless Access Authorization'),
            ('AC.3.1.17', 'Wireless Access Protection'),
            ('AC.3.1.18', 'Mobile Device Control'),
            ('AC.3.1.19', 'Mobile Device Encryption'),
            ('AC.3.1.21', 'Portable Storage Control'),
            ('AC.3.1.23', 'Network Capability Control'),
            ('AC.3.1.24', 'Device Identification'),
            ('AC.3.1.25', 'Hardware Locks'),
            ('AC.3.1.26', 'Data Execution Prevention')
        ]
    },
    'AT': {
        'name': 'Awareness and Training',
        'code': 'AT',
        'total_controls': 5,
        'l1_controls': [],
        'l2_controls': [
            ('AT.3.2.1', 'Security Awareness Training'),
            ('AT.3.2.2', 'Social Media Training'),
            ('AT.3.2.3', 'Role-based Security Training'),
            ('AT.3.2.4', 'Training Content Updates'),
            ('AT.3.2.5', 'Training Records')
        ]
    },
    'AU': {
        'name': 'Audit and Accountability',
        'code': 'AU',
        'total_controls': 9,
        'l1_controls': [],
        'l2_controls': [
            ('AU.3.3.1', 'Audit Logs'),
            ('AU.3.3.2', 'Audit Record Content'),
            ('AU.3.3.3', 'Audit Review and Analysis'),
            ('AU.3.3.4', 'Audit Reduction and Reporting'),
            ('AU.3.3.5', 'Audit Record Generation'),
            ('AU.3.3.6', 'Audit Information Protection'),
            ('AU.3.3.7', 'Near Real-time Analysis'),
            ('AU.3.3.8', 'Audit Information Protection'),
            ('AU.3.3.9', 'Audit Management')
        ]
    },
    'CA': {
        'name': 'Security Assessment',
        'code': 'CA',
        'total_controls': 8,
        'l1_controls': [],
        'l2_controls': [
            ('CA.3.12.1', 'Security Control Assessment'),
            ('CA.3.12.2', 'Plans of Action and Milestones'),
            ('CA.3.12.3', 'POA&M Development'),
            ('CA.3.12.4', 'Assessment Findings Analysis'),
            ('CA.3.12.5', 'Assessment Scope Definition'),
            ('CA.3.12.6', 'Continuous Monitoring'),
            ('CA.3.12.7', 'Monitoring Strategy'),
            ('CA.3.12.8', 'Security State Reporting')
        ]
    },
    'CM': {
        'name': 'Configuration Management',
        'code': 'CM',
        'total_controls': 11,
        'l1_controls': [],
        'l2_controls': [
            ('CM.3.4.1', 'Configuration Baselines'),
            ('CM.3.4.2', 'Security Configuration Settings'),
            ('CM.3.4.3', 'Configuration Change Control'),
            ('CM.3.4.4', 'Security Impact Analysis'),
            ('CM.3.4.5', 'Least Functionality'),
            ('CM.3.4.6', 'User-installed Software Control'),
            ('CM.3.4.7', 'Nonessential Program Control'),
            ('CM.3.4.8', 'Software Execution Policy'),
            ('CM.3.4.9', 'User-installed Hardware Control'),
            ('CM.3.4.10', 'System Logging'),
            ('CM.3.4.11', 'Configuration Updates')
        ]
    },
    'IA': {
        'name': 'Identification and Authentication',
        'code': 'IA',
        'total_controls': 11,
        'l1_controls': [
            ('IA.L1-3.5.1', 'User Identification'),
            ('IA.L1-3.5.2', 'User Authentication')
        ],
        'l2_controls': [
            ('IA.3.5.3', 'Multi-factor Authentication'),
            ('IA.3.5.4', 'Replay-resistant Authentication'),
            ('IA.3.5.5', 'Identifier Reuse Prevention'),
            ('IA.3.5.6', 'Identifier Management'),
            ('IA.3.5.7', 'Password Complexity'),
            ('IA.3.5.8', 'Password Reuse'),
            ('IA.3.5.9', 'Temporary Passwords'),
            ('IA.3.5.10', 'Cryptographic Password Protection'),
            ('IA.3.5.11', 'Authentication Feedback')
        ]
    },
    'IR': {
        'name': 'Incident Response',
        'code': 'IR',
        'total_controls': 3,
        'l1_controls': [],
        'l2_controls': [
            ('IR.3.6.1', 'Incident Handling Capability'),
            ('IR.3.6.2', 'Incident Tracking and Reporting'),
            ('IR.3.6.3', 'Incident Response Testing')
        ]
    },
    'MA': {
        'name': 'Maintenance',
        'code': 'MA',
        'total_controls': 6,
        'l1_controls': [],
        'l2_controls': [
            ('MA.3.7.1', 'System Maintenance'),
            ('MA.3.7.2', 'CUI Maintenance Protection'),
            ('MA.3.7.3', 'Equipment Sanitization'),
            ('MA.3.7.4', 'Maintenance Records'),
            ('MA.3.7.5', 'Remote Maintenance Authentication'),
            ('MA.3.7.6', 'Maintenance Personnel Supervision')
        ]
    },
    'MP': {
        'name': 'Media Protection',
        'code': 'MP',
        'total_controls': 8,
        'l1_controls': [
            ('MP.L1-3.8.3', 'Media Sanitization')
        ],
        'l2_controls': [
            ('MP.3.8.1', 'CUI Media Identification'),
            ('MP.3.8.2', 'Media Access Control'),
            ('MP.3.8.4', 'Media Distribution Marking'),
            ('MP.3.8.5', 'CUI Media Access Control'),
            ('MP.3.8.6', 'Mobile Device Restrictions'),
            ('MP.3.8.7', 'Portable Storage Control'),
            ('MP.3.8.8', 'Media Transport Protection')
        ]
    },
    'PE': {
        'name': 'Physical Protection',
        'code': 'PE',
        'total_controls': 8,
        'l1_controls': [
            ('PE.L1-3.10.1', 'Physical Access Authorization'),
            ('PE.L1-3.10.2', 'Physical Facility Protection'),
            ('PE.L1-3.10.3', 'Physical Access Device Control'),
            ('PE.L1-3.10.4', 'Visitor Access and Monitoring')
        ],
        'l2_controls': [
            ('PE.3.10.5', 'Physical System Access Control'),
            ('PE.3.10.6', 'Physical Access Audit Logs'),
            ('PE.3.10.7', 'Alternate Worksite Safeguarding'),
            ('PE.3.10.8', 'Badge-based Access Control')
        ]
    },
    'PS': {
        'name': 'Personnel Security',
        'code': 'PS',
        'total_controls': 3,
        'l1_controls': [],
        'l2_controls': [
            ('PS.3.9.1', 'Personnel Screening'),
            ('PS.3.9.2', 'Personnel Action Protection'),
            ('PS.3.9.3', 'Personnel Termination')
        ]
    },
    'RA': {
        'name': 'Risk Assessment',
        'code': 'RA',
        'total_controls': 6,
        'l1_controls': [],
        'l2_controls': [
            ('RA.3.11.1', 'Risk Assessment'),
            ('RA.3.11.2', 'Vulnerability Scanning'),
            ('RA.3.11.3', 'Vulnerability Remediation'),
            ('RA.3.11.4', 'Risk Assessment Updates'),
            ('RA.3.11.5', 'Threat Intelligence Sharing'),
            ('RA.3.11.6', 'Threat Intelligence Analysis')
        ]
    },
    'SC': {
        'name': 'System and Communications Protection',
        'code': 'SC',
        'total_controls': 19,
        'l1_controls': [
            ('SC.L1-3.13.1', 'Communications Monitoring'),
            ('SC.L1-3.13.5', 'License Restrictions')
        ],
        'l2_controls': [
            ('SC.3.13.2', 'System Management Separation'),
            ('SC.3.13.3', 'Remote Device Connections'),
            ('SC.3.13.4', 'Cryptographic Mechanisms'),
            ('SC.3.13.5', 'Public Access Subnetworks'),
            ('SC.3.13.6', 'Network Communications Denial'),
            ('SC.3.13.7', 'Split Tunneling Prevention'),
            ('SC.3.13.8', 'CUI Transmission Protection'),
            ('SC.3.13.9', 'Session Authenticity Protection'),
            ('SC.3.13.10', 'Cryptographic Key Management'),
            ('SC.3.13.11', 'FIPS-validated Cryptography'),
            ('SC.3.13.12', 'Remote Device Activation'),
            ('SC.3.13.13', 'Mobile Code Control'),
            ('SC.3.13.14', 'Voice over IP Control'),
            ('SC.3.13.15', 'System Media Protection'),
            ('SC.3.13.16', 'CUI Transmission Protection'),
            ('SC.3.13.17', 'Public System CUI Control')
        ]
    },
    'SI': {
        'name': 'System and Information Integrity',
        'code': 'SI',
        'total_controls': 10,
        'l1_controls': [
            ('SI.L1-3.14.1', 'Information System Flaws'),
            ('SI.L1-3.14.2', 'Malicious Code Protection'),
            ('SI.L1-3.14.3', 'Malicious Code Updates'),
            ('SI.L1-3.14.4', 'Malicious Code Scanning')
        ],
        'l2_controls': [
            ('SI.3.14.5', 'Security Alerts and Advisories'),
            ('SI.3.14.6', 'Attack Indicators Monitoring'),
            ('SI.3.14.7', 'Unauthorized Change Detection'),
            ('SI.3.14.8', 'Software Whitelisting'),
            ('SI.3.14.9', 'Software Blacklisting'),
            ('SI.3.14.10', 'Malicious Code Protection Updates')
        ]
    }
}

def create_policy_document(family_code, family_data):
    """Create a CMMC Level 2 policy document in Capital Cyber format"""
    
    doc = Document()
    
    # Set up document properties
    title = f"{family_data['name']} ({family_code}) Policies and Procedures"
    
    # Title
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Empty line
    
    # Document Revision History Table
    doc.add_heading('Document Revision History', level=2)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Version'
    hdr_cells[2].text = 'Description'
    hdr_cells[3].text = 'Author'
    
    # Data row
    row_cells = table.rows[1].cells
    row_cells[0].text = 'March 1, 2026'
    row_cells[1].text = '1.0'
    row_cells[2].text = 'Initial Document'
    row_cells[3].text = 'Capital Cyber'
    
    doc.add_paragraph()  # Empty line
    
    # Approved By section
    doc.add_heading('Approved By', level=2)
    approval_table = doc.add_table(rows=2, cols=4)
    approval_table.style = 'Table Grid'
    
    # Header row
    approval_hdr = approval_table.rows[0].cells
    approval_hdr[0].text = 'Name'
    approval_hdr[1].text = 'Title'
    approval_hdr[2].text = 'Date'
    approval_hdr[3].text = 'Signature'
    
    # Data row
    approval_row = approval_table.rows[1].cells
    approval_row[0].text = '[Senior Official Name]'
    approval_row[1].text = '[Title]'
    approval_row[2].text = ''
    approval_row[3].text = ''
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction',
        '2. Purpose', 
        '3. Scope',
        '4. Roles and Responsibilities',
        '5. Management Commitment',
        '6. Authority and References',
        '7. Compliance',
        '8. Procedural Requirements'
    ]
    
    # Add L1 controls to TOC if any
    section_num = 9
    if family_data['l1_controls']:
        for control_code, control_name in family_data['l1_controls']:
            toc_items.append(f'{section_num}. {control_code}: {control_name}')
            section_num += 1
    
    # Add L2 controls to TOC
    for control_code, control_name in family_data['l2_controls']:
        toc_items.append(f'{section_num}. {control_code}: {control_name}')
        section_num += 1
    
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    intro_text = f"""[COMPANY NAME] is a precision sheet metal fabrication, CNC machining, and CNC fiber laser company founded in 1968, serving defense and commercial clients including GE. As a handler of Federal Contract Information (FCI) and Controlled Unclassified Information (CUI), [COMPANY NAME] must comply with the Cybersecurity Maturity Model Certification (CMMC) framework. This document establishes the {family_data['name']} policies and procedures aligned with NIST SP 800-171 Revision 2 to ensure that {family_data['name'].lower()} requirements are properly implemented and maintained."""
    doc.add_paragraph(intro_text)
    
    # 2. Purpose
    doc.add_heading('2. Purpose', level=1)
    purpose_text = f"""The purpose of this policy is to define the requirements for implementing {family_data['name'].lower()} controls within [COMPANY NAME] information systems, networks, and data. This policy ensures that all {family_data['name'].lower()} requirements are met in accordance with NIST SP 800-171 Revision 2 and CMMC Level 2 standards."""
    doc.add_paragraph(purpose_text)
    
    # 3. Scope
    doc.add_heading('3. Scope', level=1)
    scope_text = """This policy applies to all [COMPANY NAME] employees, contractors, subcontractors, and third-party personnel who access, process, store, or transmit organizational information and information systems. This includes all systems and data environments that handle Federal Contract Information (FCI), Controlled Unclassified Information (CUI), Covered Defense Information (CDI), and Export Controlled data."""
    doc.add_paragraph(scope_text)
    
    # 4. Roles and Responsibilities
    doc.add_heading('4. Roles and Responsibilities', level=1)
    roles_text = """The following individuals and groups are responsible for implementing and maintaining this policy:"""
    doc.add_paragraph(roles_text)
    
    roles_table = doc.add_table(rows=7, cols=2)
    roles_table.style = 'Table Grid'
    
    roles_data = [
        ('Individual/Role', 'Responsibility'),
        ('[Senior Official Name]', f'Overall accountability for {family_data["name"].lower()} policy implementation and compliance'),
        ('[To Be Assigned]', f'Day-to-day management of {family_data["name"].lower()} controls and procedures'),
        ('[MSP Name]', f'Technical implementation and maintenance of {family_data["name"].lower()} systems'),
        ('[MSSP Name]', f'Monitoring and incident response for {family_data["name"].lower()} events'),
        ('[MCP Name]', f'Third-party compliance validation for {family_data["name"].lower()} controls'),
        ('All Employees', f'Adherence to {family_data["name"].lower()} policies and procedures in daily operations')
    ]
    
    for i, (role, responsibility) in enumerate(roles_data):
        roles_table.rows[i].cells[0].text = role
        roles_table.rows[i].cells[1].text = responsibility
    
    # 5. Management Commitment
    doc.add_heading('5. Management Commitment', level=1)
    commitment_text = f"""[COMPANY NAME] leadership is committed to implementing and maintaining effective {family_data['name'].lower()} controls. Senior management provides the necessary resources, oversight, and authority to ensure compliance with all applicable regulations and standards. Regular reviews and assessments are conducted to verify the effectiveness of implemented controls."""
    doc.add_paragraph(commitment_text)
    
    # 6. Authority and References
    doc.add_heading('6. Authority and References', level=1)
    references_text = """This policy is established under the authority of [Senior Official Name] and is based on the following references:
    
• NIST Special Publication 800-171 Revision 2
• CMMC Model Version 2.0
• 32 CFR Part 170 (CMMC Final Rule)
• DFARS 252.204-7012
• FAR 52.204-21
• FIPS 140-2/3
• FIPS 199
• FIPS 200"""
    doc.add_paragraph(references_text)
    
    # 7. Compliance
    doc.add_heading('7. Compliance', level=1)
    compliance_text = f"""All personnel must comply with this {family_data['name']} policy. Failure to comply may result in disciplinary action, up to and including termination of employment or contract. Regular training and awareness programs ensure all personnel understand their responsibilities under this policy."""
    doc.add_paragraph(compliance_text)
    
    # 8. Procedural Requirements
    doc.add_heading('8. Procedural Requirements', level=1)
    proc_text = f"""The following procedural requirements apply to all {family_data['name'].lower()} activities:
    
• All controls must be implemented according to organizational procedures
• Regular monitoring and assessment of control effectiveness
• Documentation of all {family_data['name'].lower()} activities and decisions
• Prompt reporting of any control failures or incidents
• Continuous improvement based on lessons learned and threat intelligence
• Implementation details to be customized per client requirements and environment"""
    doc.add_paragraph(proc_text)
    
    # Add control sections
    section_num = 9
    
    # Add L1 controls if any
    if family_data['l1_controls']:
        for control_code, control_name in family_data['l1_controls']:
            doc.add_heading(f'{section_num}. {control_code}: {control_name}', level=1)
            
            # Requirement subsection
            doc.add_heading('Requirement', level=2)
            req_text = f"""[COMPANY NAME] shall implement {control_name.lower()} controls in accordance with NIST SP 800-171 Revision 2 requirements. Implementation details to be customized per client requirements and operational environment."""
            doc.add_paragraph(req_text)
            
            # Implementation subsection  
            doc.add_heading('Implementation', level=2)
            impl_text = f"""[COMPANY NAME] implements {control_name.lower()} through a combination of technical controls, administrative procedures, and monitoring activities. Specific implementation details are documented in system-specific security plans and standard operating procedures. [MSP Name], [MSSP Name], and [MCP Name] provide support for implementation and ongoing maintenance of these controls."""
            doc.add_paragraph(impl_text)
            
            section_num += 1
    
    # Add L2 controls
    for control_code, control_name in family_data['l2_controls']:
        doc.add_heading(f'{section_num}. {control_code}: {control_name}', level=1)
        
        # Requirement subsection
        doc.add_heading('Requirement', level=2)
        req_text = f"""[COMPANY NAME] shall implement {control_name.lower()} controls in accordance with NIST SP 800-171 Revision 2 requirements. Implementation details to be customized per client requirements and operational environment."""
        doc.add_paragraph(req_text)
        
        # Implementation subsection
        doc.add_heading('Implementation', level=2)
        impl_text = f"""[COMPANY NAME] implements {control_name.lower()} through a combination of technical controls, administrative procedures, and monitoring activities. Specific implementation details are documented in system-specific security plans and standard operating procedures. [MSP Name], [MSSP Name], and [MCP Name] provide support for implementation and ongoing maintenance of these controls."""
        doc.add_paragraph(impl_text)
        
        section_num += 1
    
    return doc

def main():
    """Generate all 14 CMMC Level 2 policy documents"""
    output_dir = '/root/.openclaw/workspace/skills/cmmc-grc/assets/templates/L2-capital-cyber'
    
    print("Generating CMMC Level 2 Policy Documents in Capital Cyber format...")
    print(f"Output directory: {output_dir}")
    print()
    
    for family_code, family_data in CONTROL_FAMILIES.items():
        print(f"Creating {family_code} - {family_data['name']} ({family_data['total_controls']} controls)")
        
        # Create the document
        doc = create_policy_document(family_code, family_data)
        
        # Save the document
        filename = f"{family_code}-{family_data['name'].lower().replace(' ', '-').replace('&', 'and')}-policy.docx"
        filepath = os.path.join(output_dir, filename)
        
        doc.save(filepath)
        print(f"  Saved: {filename}")
    
    print()
    print("All 14 CMMC Level 2 policy documents have been generated successfully!")
    print(f"Files saved to: {output_dir}")

if __name__ == '__main__':
    main()